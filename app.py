import io
import os
import re
import shutil
from datetime import datetime

import streamlit as st
from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
from docx import Document

# =====================================================
# PAGE CONFIG — MUST BE FIRST STREAMLIT COMMAND
# =====================================================
st.set_page_config(page_title="PDF → Word (OCR)", page_icon="🧾", layout="wide")

# =====================================================
# BRANDING (Chumcred Limited Logo)
# =====================================================
LOGO_PATH = os.path.join("assets", "chumcred_logo.png")

if os.path.exists(LOGO_PATH):
    col1, col2 = st.columns([1, 5])
    with col1:
        st.image(LOGO_PATH, width=90)
    with col2:
        st.markdown("## Chumcred Limited")
        st.caption("PDF → Word Converter")

    with st.sidebar:
        st.image(LOGO_PATH, width=200)
else:
    st.warning("Logo not found. Please add: assets/chumcred_logo.png")

# =====================================================
# OCR ENGINE CONFIG (Windows + Streamlit Cloud safe)
# =====================================================
POPPLER_BIN = r"C:\Users\ADEGBIE ADEKUNLE\poppler\poppler-25.12.0\Library\bin"
TESSERACT_EXE_WIN = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# --- Poppler (Windows only) ---
if os.name == "nt" and os.path.isdir(POPPLER_BIN):
    os.environ["PATH"] = POPPLER_BIN + os.pathsep + os.environ.get("PATH", "")

# --- Tesseract (Windows OR Linux) ---
if os.name == "nt":
    # Windows
    if os.path.exists(TESSERACT_EXE_WIN):
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_EXE_WIN
        os.environ["PATH"] = r"C:\Program Files\Tesseract-OCR" + os.pathsep + os.environ.get("PATH", "")
else:
    # Linux (Streamlit Cloud)
    tesseract_path = shutil.which("tesseract")
    if tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path

# -----------------------------
# Helpers
# -----------------------------
def sanitize_filename(name: str) -> str:
    name = re.sub(r"[^\w\-. ]+", "_", name).strip()
    return name or "output"


def ocr_image(img: Image.Image, lang: str, psm: int = 3) -> str:
    config = f"--oem 3 --psm {psm}"
    return pytesseract.image_to_string(img, lang=lang, config=config)


def build_docx_from_text(pages_text: list[str], title: str = None) -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for i, txt in enumerate(pages_text, start=1):
        doc.add_heading(f"Page {i}", level=2)
        for para in [p.strip() for p in txt.split("\n\n") if p.strip()]:
            doc.add_paragraph(para)
        if i != len(pages_text):
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# -----------------------------
# Streamlit UI
# -----------------------------
st.title("🧾 PDF → Word Converter (OCR)")
st.write(
    "Upload a PDF (including scanned PDFs). The app will run OCR on each page and generate a Word (.docx) file."
)

with st.sidebar:
    st.header("OCR Settings")

    lang = st.selectbox("OCR language", ["eng", "fra", "deu", "spa", "ita", "por"], index=0)

    psm = st.selectbox(
        "Page segmentation mode (PSM)",
        options=[3, 4, 6, 11, 12],
        index=0,
        help=(
            "3 = automatic (recommended). 6 = assume a single uniform block of text. "
            "11/12 = sparse text."
        ),
    )

    dpi = st.slider(
        "Render DPI (higher = better OCR, slower)",
        min_value=150,
        max_value=400,
        value=250,
        step=50,
    )

    st.caption(
        "Tip: For very blurry scans, try higher DPI and PSM=6. "
        "For receipts or scattered text, try PSM=11 or 12."
    )

uploaded = st.file_uploader("Upload a PDF", type=["pdf"])

if uploaded:
    base_name = sanitize_filename(os.path.splitext(uploaded.name)[0])
    st.info(f"File: **{uploaded.name}**")

    run = st.button("🚀 Convert to Word", type="primary")

    if run:
        pdf_bytes = uploaded.read()

        try:
            with st.spinner("Rendering PDF pages..."):
                kwargs = dict(dpi=dpi)

                # Only pass poppler_path on Windows
                if os.name == "nt" and os.path.isdir(POPPLER_BIN):
                    kwargs["poppler_path"] = POPPLER_BIN

                pages = convert_from_bytes(pdf_bytes, **kwargs)

            st.success(f"Rendered {len(pages)} page(s).")

            pages_text = []
            progress = st.progress(0)
            status = st.empty()

            for idx, page_img in enumerate(pages, start=1):
                status.write(f"OCR in progress: page {idx}/{len(pages)} ...")

                page_img = page_img.convert("RGB")
                text = ocr_image(page_img, lang=lang, psm=psm)
                pages_text.append(text)

                progress.progress(int(idx / len(pages) * 100))

            status.write("Building Word document...")
            title = f"OCR Output - {uploaded.name}"
            docx_bytes = build_docx_from_text(pages_text, title=title)

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            out_name = f"{base_name}_OCR_{ts}.docx"

            st.success("✅ Done! Download your Word document below.")
            st.download_button(
                label="⬇️ Download Word (.docx)",
                data=docx_bytes,
                file_name=out_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            with st.expander("Preview extracted text (first 2 pages)"):
                for i, txt in enumerate(pages_text[:2], start=1):
                    st.subheader(f"Page {i}")
                    st.text(txt[:4000] + ("..." if len(txt) > 4000 else ""))

        except Exception as e:
            st.error("Conversion failed.")
            st.exception(e)
else:
    st.warning("Upload a PDF to begin.")
